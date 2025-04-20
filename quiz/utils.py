# quiz/utils.py
import docx
import pandas as pd
from .models import Quiz, Question, Option, Quiz_Question
from django.utils.text import slugify
from django.utils import timezone
from institutions.models import Institution  

def get_default_institution(user):
    return Institution.objects.filter(user=user).first()  # assuming you have a relation

def parse_word_file(file_path, user, title,description, is_public, start_time, end_time, shuffle_questions, price):
    document = docx.Document(file_path)
    institution = get_default_institution(user)

    quiz = Quiz.objects.create(
        title=title,
        created_by=user,
        institution=institution,
        is_public=is_public,
        start_time=start_time,
        end_time=end_time,
        shuffle_questions = shuffle_questions,
        price = price,
        description = description, 
    )

    current_question = ""
    options = []
    correct_option_key = ""

    option_keys = ['options_A', 'options_B', 'options_C', 'options_D']
    current_option_index = 0

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if text.endswith('?'):
            # Save previous question if any
            if current_question and options:
                Quiz_Question.objects.create(
                    quiz=quiz,
                    question=current_question,
                    options_A=options[0] if len(options) > 0 else '',
                    options_B=options[1] if len(options) > 1 else '',
                    options_C=options[2] if len(options) > 2 else '',
                    options_D=options[3] if len(options) > 3 else '',
                    is_correct=correct_option_key
                )

            # Reset for new question
            current_question = text
            options = []
            correct_option_key = ""
            current_option_index = 0

        elif current_question and current_option_index < 4:
            is_correct = text.startswith('*')
            option_text = text[1:].strip() if is_correct else text
            options.append(option_text)

            if is_correct:
                correct_option_key = option_keys[current_option_index]

            current_option_index += 1

    # Handle last question
    if current_question and options:
        Quiz_Question.objects.create(
            quiz=quiz,
            question=current_question,
            options_A=options[0] if len(options) > 0 else '',
            options_B=options[1] if len(options) > 1 else '',
            options_C=options[2] if len(options) > 2 else '',
            options_D=options[3] if len(options) > 3 else '',
            is_correct=correct_option_key
        )


def parse_excel_file(file_path, user, title, description, is_public, start_time, end_time, shuffle_questions, price):
    df = pd.read_excel(file_path)
    institution = get_default_institution(user)

    quiz = Quiz.objects.create(
        title=title,
        created_by=user,
        institution=institution,
        is_public=is_public,
        start_time=start_time,
        end_time=end_time,
        shuffle_questions = shuffle_questions,
        price = price,
        description = description,
    )

    for _, row in df.iterrows():
        question_text = row.get('Question')
        if pd.notna(question_text):
            # Fetch options
            option_A = row.get('Option_A') or ''
            option_B = row.get('Option_B') or ''
            option_C = row.get('Option_C') or ''
            option_D = row.get('Option_D') or ''

            correct_col = str(row.get('Correct')).strip().upper()
            correct_option_map = {
                'A': 'options_A',
                'B': 'options_B',
                'C': 'options_C',
                'D': 'options_D',
            }
            correct_key = correct_option_map.get(correct_col, '')

            # Create question
            Quiz_Question.objects.create(
                quiz=quiz,
                question=question_text,
                options_A=option_A,
                options_B=option_B,
                options_C=option_C,
                options_D=option_D,
                is_correct=correct_key
            )

# Question | Option_A | Option_B | Option_C | Option_D | Correct


# def parse_word_file(file_path, user):
#     document = docx.Document(file_path)
#     institution = get_default_institution(user)
#     quiz = Quiz.objects.create(title=f"Word Quiz {user.username}", 
#         created_by=user,
#         institution = institution,
#         is_public = True,
#         start_time=timezone.now(),
#         end_time=timezone.now() + timezone.timedelta(hours=1))
    
#     current_question = None

#     for para in document.paragraphs:
#         text = para.text.strip()
#         if text:
#             if text.endswith('?'):
#                 current_question = Question.objects.create(quiz=quiz, text=text)
#             elif current_question:
#                 is_correct = text.startswith('*')
#                 option_text = text[1:].strip() if is_correct else text
#                 Option.objects.create(question=current_question, text=option_text, is_correct=is_correct)

# def parse_excel_file(file_path, user):
#     df = pd.read_excel(file_path)
#     institution = get_default_institution(user)

#     quiz = Quiz.objects.create(
#         title=f"Excel Quiz {user.username}",
#         created_by=user,
#         institution=institution,
#         is_public=True,
#         start_time=timezone.now(),
#         end_time=timezone.now() + timezone.timedelta(hours=1)
#     )

#     for _, row in df.iterrows():
#         question_text = row.get('Question')
#         if pd.notna(question_text):
#             question = Question.objects.create(quiz=quiz, text=question_text)
#             for i in ['A', 'B', 'C', 'D']:
#                 option_col = f"Option_{i}"
#                 correct_col = 'Correct'
#                 option_text = row.get(option_col)
#                 if pd.notna(option_text):
#                     is_correct = str(row.get(correct_col)).strip().upper() == i
#                     Option.objects.create(
#                         question=question,
#                         text=option_text,
#                         is_correct=is_correct
#                     )
